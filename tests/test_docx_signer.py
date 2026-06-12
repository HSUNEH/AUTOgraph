from pathlib import Path

from docx import Document
from PIL import Image

from autograph.signers.docx import sign_docx


def _signature_png(path: Path) -> None:
    Image.new("RGBA", (240, 80), (0, 0, 0, 0)).save(path)


def test_sign_docx_replaces_named_placeholder_with_signature_image(tmp_path: Path):
    source = tmp_path / "contract.docx"
    signature = tmp_path / "alice.png"
    output = tmp_path / "signed.docx"

    _signature_png(signature)
    doc = Document()
    doc.add_paragraph("Agreement")
    doc.add_paragraph("Signer: {{signature: Alice Kim}}")
    doc.save(source)

    result = sign_docx(source, output, signature, signer="Alice Kim")

    signed = Document(output)
    assert result.signed is True
    assert result.placements[0].signer == "Alice Kim"
    assert "{{signature: Alice Kim}}" not in "\n".join(p.text for p in signed.paragraphs)
    assert len(signed.inline_shapes) == 1


def test_sign_docx_replaces_placeholder_split_across_runs(tmp_path: Path):
    source = tmp_path / "contract.docx"
    signature = tmp_path / "alice.png"
    output = tmp_path / "signed.docx"

    _signature_png(signature)
    doc = Document()
    paragraph = doc.add_paragraph("Signer: ")
    paragraph.add_run("{{signature")
    paragraph.add_run(": Alice Kim}}")
    doc.save(source)

    result = sign_docx(source, output, signature, signer="Alice Kim")

    signed = Document(output)
    assert result.signed is True
    assert "{{signature: Alice Kim}}" not in "\n".join(p.text for p in signed.paragraphs)
    assert len(signed.inline_shapes) == 1
