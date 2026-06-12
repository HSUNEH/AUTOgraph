from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader

SUPPORTED_TYPES = {"pdf", "docx", "doc", "hwp", "hwpx"}
MAX_HWPX_XML_MEMBERS = 1_000
MAX_HWPX_XML_MEMBER_BYTES = 2_000_000
MAX_HWPX_TOTAL_XML_BYTES = 20_000_000
MAX_HWPX_COMPRESSION_RATIO = 100


def sniff_document_type(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower().lstrip(".")
    if suffix in SUPPORTED_TYPES:
        return suffix
    raise ValueError(f"Unsupported document type: {path.suffix or '(none)'}")


def extract_text(path: str | Path) -> str:
    path = Path(path)
    kind = sniff_document_type(path)
    if kind == "pdf":
        return extract_pdf_text(path)
    if kind == "docx":
        return extract_docx_text(path)
    if kind == "hwpx":
        return extract_hwpx_text(path)
    if kind == "hwp":
        return extract_hwp_text(path)
    raise NotImplementedError(
        "Legacy .doc text extraction requires external conversion to .docx or .pdf."
    )


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(chunks)


def extract_docx_text(path: Path) -> str:
    document = Document(str(path))
    chunks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def extract_hwpx_text(path: Path) -> str:
    chunks: list[str] = []
    xml_members = 0
    total_xml_bytes = 0
    with zipfile.ZipFile(path) as archive:
        for info in archive.infolist():
            name = info.filename
            if not name.lower().endswith(".xml"):
                continue
            if not (name.startswith("Contents/") or name.startswith("Preview/")):
                continue
            xml_members += 1
            total_xml_bytes += info.file_size
            compressed_size = max(info.compress_size, 1)
            if xml_members > MAX_HWPX_XML_MEMBERS:
                raise ValueError("HWPX archive has too many XML members")
            if info.file_size > MAX_HWPX_XML_MEMBER_BYTES:
                raise ValueError(f"HWPX XML member is too large: {name}")
            if total_xml_bytes > MAX_HWPX_TOTAL_XML_BYTES:
                raise ValueError("HWPX archive XML payload is too large")
            if info.file_size / compressed_size > MAX_HWPX_COMPRESSION_RATIO:
                raise ValueError(f"HWPX XML member compression ratio is too large: {name}")

            data = archive.read(info)
            try:
                root = ET.fromstring(data)
            except ET.ParseError:
                continue
            for element in root.iter():
                if element.text and element.text.strip():
                    chunks.append(element.text.strip())
    return "\n".join(chunks)


def extract_hwp_text(path: Path) -> str:
    try:
        from hwp_hwpx_parser import HwpParser  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "HWP extraction needs the optional dependency: pip install autograph-agent[hwp]"
        ) from exc

    parser = HwpParser(str(path))
    return str(parser.extract_text())
