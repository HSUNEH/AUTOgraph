from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches

from autograph.detectors import find_signature_requests
from autograph.models import SignaturePlacement, SignResult
from autograph.parsers import extract_docx_text


def sign_docx(
    input_path: str | Path,
    output_path: str | Path,
    signature_path: str | Path,
    *,
    signer: str,
    width_inches: float = 1.5,
) -> SignResult:
    input_path = Path(input_path)
    output_path = Path(output_path)
    signature_path = Path(signature_path)

    text = extract_docx_text(input_path)
    requests = find_signature_requests(text)
    request = next((item for item in requests if item.signer.casefold() == signer.casefold()), None)
    if request is None:
        shutil.copyfile(input_path, output_path)
        return SignResult(
            input_path=input_path,
            output_path=output_path,
            signed=False,
            format="docx",
            warnings=[f"No signature placeholder found for signer: {signer}"],
        )

    document = Document(str(input_path))
    placement: SignaturePlacement | None = None
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        if request.placeholder not in paragraph.text:
            continue

        prefix, suffix = paragraph.text.split(request.placeholder, 1)
        for run in paragraph.runs:
            run.text = ""
        target_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        target_run.text = prefix
        target_run.add_picture(str(signature_path), width=Inches(width_inches))
        if suffix:
            paragraph.add_run(suffix)
        placement = SignaturePlacement(
            signer=signer,
            placeholder=request.placeholder,
            page=None,
            x=float(paragraph_index),
            y=0.0,
            width=width_inches,
            height=0.0,
        )
        break

    if placement is None:
        shutil.copyfile(input_path, output_path)
        return SignResult(
            input_path=input_path,
            output_path=output_path,
            signed=False,
            format="docx",
            warnings=["Placeholder was detected in text but not in editable paragraph runs."],
        )

    document.save(str(output_path))
    return SignResult(
        input_path=input_path,
        output_path=output_path,
        signed=True,
        format="docx",
        placements=[placement],
    )
